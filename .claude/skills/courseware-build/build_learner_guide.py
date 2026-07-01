#!/usr/bin/env python3
"""Single-source Learner Guide generator for WSQ TGS-2021010366.

Emits BOTH:
  - LG-Application-Integration-with-Docker-and-Kubernetes.md   (repo root)
  - courseware/LG-Application-Integration-with-Docker-and-Kubernetes.docx
from labs_data.py, so the slide deck, labs, lesson plan, LG and MD all stay aligned.
"""
import os, re, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import prodoc
import labs_data as L
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.environ.get("COURSE_REPO") or os.getcwd()
TITLE = "Application Integration with Docker and Kubernetes"
CODE = "TGS-2021010366"
VERSION = "1.0"
ASSETS = os.path.join(ROOT, "courseware", "assets")
VERSIONS = [
    ("1.0", "30 June 2026",
     "First version — step-by-step guide to all 19 Docker & Kubernetes labs, built around "
     "the TaskBoard sample app; MD and DOCX generated from one source",
     "Tertiary Infotech Academy Pte Ltd"),
]

# ---------------------------------------------------------------- assemble blocks
B = []
def h1(t): B.append(("h1", t))
def h2(t): B.append(("h2", t))
def h3(t): B.append(("h3", t))
def p(t):  B.append(("p", t))
def bullets(xs): B.append(("bullets", list(xs)))
def code(t): B.append(("code", t))
def table(rs): B.append(("table", rs))
def note(t): B.append(("note", t))
def test(t): B.append(("test", t))
def rule(): B.append(("rule",))

h1("Application Integration with Docker and Kubernetes — Step-by-Step Learner Guide")
p("Welcome! This guide walks you command-by-command through every hands-on lab in the WSQ "
  "course **Application Integration with Docker and Kubernetes** (Course Code: TGS-2021010366). "
  "Over two days you build and deploy one real application — **TaskBoard**, a Flask task "
  "tracker — from a single container, to a multi-service Docker Compose stack, to a scalable "
  "Kubernetes deployment.")
p("Work through the labs in order: each one builds on the last. Whenever you see a "
  "**Test it** box, stop and confirm the result before moving on. All labs also run in the "
  "browser on KillerCoda — the link is at the top of each lab.")
note("Course flow at a glance — **Day 1 (Docker):** commands, Dockerfile & build, CMD vs "
     "ENTRYPOINT, volumes, networking, environment config, Docker Hub, Docker Compose "
     "(Labs 1–12). **Day 2 (Kubernetes):** Pods, Namespaces, Deployments, rollouts, "
     "Services, storage, Jobs & CronJobs (Labs 13–19).")

# ---- 0. setup
rule()
h2("0. Before You Start — Setup & Prerequisites")
h3("0.1 What you need")
table([
    ["Tool", "Used for", "Where to get it"],
    ["Docker Engine / Desktop", "Building & running containers (all Day 1 labs)", "docker.com (or use KillerCoda in the browser)"],
    ["Docker Compose", "Multi-service labs 10–12", "Included with Docker Desktop / Docker Engine"],
    ["kubectl + a cluster", "All Day 2 Kubernetes labs", "KillerCoda playground, or minikube/kind locally"],
    ["Docker Hub account", "Pushing your image (Lab 9)", "hub.docker.com — free"],
])
h3("0.2 Two ways to run every lab")
p("**Option A — KillerCoda (fastest).** Each lab header has a KillerCoda link; the commands "
  "run in a browser terminal with Docker and Kubernetes pre-installed — nothing to install.")
p("**Option B — Local.** Install Docker Desktop (gives you Docker + Compose) and a local "
  "Kubernetes (minikube, kind, or Docker Desktop's built-in Kubernetes) for Day 2.")
h3("0.3 Get the lab files")
p("Every lab folder under **`labs/`** is self-contained — it holds the `lab.md` steps plus "
  "the working files (app code, `Dockerfile`, `docker-compose.yml`, Kubernetes YAML). The "
  "shared sample app lives in **`labs/app/`**.")
note("**GitHub repo:** https://github.com/tertiarycourses/TGS-2021010366-Application-Integration-with-Docker-and-Kubernetes "
     " · clone it or use **Code → Download ZIP**, then `cd` into each lab folder as you go.")

# ---- per-lab sections (the single source of truth)
for lab in L.LABS:
    rule()
    h2(f"Lab {lab['num']} — {lab['title']}")
    p(f"**Folder:** `labs/{lab['slug']}/`  ·  **KillerCoda:** {lab['killercoda']}")
    h3("Goal")
    p(lab["goal"])
    h3("What you'll build")
    p(lab["build"])
    B.extend(lab["body"])      # the lab's own h3/p/steps/code/note/table blocks
    test(lab["test"])

# ---- troubleshooting + glossary
rule()
h2("Troubleshooting Cheat-Sheet")
table([
    ["Symptom", "Likely cause & fix"],
    ["`port is already allocated`", "Another container uses that host port — change `-p 8081:5000` or `docker rm -f` the other one."],
    ["`Cannot connect to the Docker daemon`", "Docker isn't running — start Docker Desktop (or use KillerCoda)."],
    ["`denied: requested access to the resource is denied` on push", "Run `docker login` and tag as `<your-user>/taskboard:1.0` (Lab 9)."],
    ["Web app can't reach Redis/DB by name", "Containers must share a **custom** network (Lab 7) or be in the same Compose file; the host is the **service/container name**."],
    ["Pod stuck `Pending`", "`kubectl describe pod <name>` — usually no schedulable node or an unbound PVC."],
    ["PVC stuck `Pending`", "No PV matches the request — check `storage` size and `accessModes` (Lab 18)."],
    ["Service returns nothing", "The Service `selector` must match the Pod `labels`; check `kubectl get endpoints`."],
])
rule()
h2("Glossary")
table([
    ["Term", "Meaning"],
    ["Image", "A read-only package with everything needed to run an app; built from a Dockerfile."],
    ["Container", "A running (or stopped) instance of an image — an isolated process."],
    ["Dockerfile", "The text recipe of instructions Docker uses to build an image."],
    ["Layer", "One filesystem change in an image; layers are cached and reused across builds."],
    ["CMD / ENTRYPOINT", "Default command vs fixed executable for a container (Lab 5)."],
    ["Named volume / Bind mount", "Docker-managed persistent storage vs a mounted host folder (Lab 6)."],
    ["Bridge network", "A virtual network; a custom bridge adds DNS so containers reach each other by name (Lab 7)."],
    ["Registry / Docker Hub", "A store of images that clusters and users pull from (Lab 9)."],
    ["Docker Compose", "A tool to define and run multi-service apps from one YAML file (Labs 10–12)."],
    ["Pod", "The smallest Kubernetes unit — one or more containers sharing network & storage."],
    ["Deployment / ReplicaSet", "Manages a desired number of identical Pods; self-heals and updates them."],
    ["Service (ClusterIP / NodePort)", "A stable network endpoint load-balancing across Pods (Lab 17)."],
    ["Namespace", "A virtual cluster for isolating environments (Lab 14)."],
    ["PV / PVC", "PersistentVolume (storage) and PersistentVolumeClaim (a request for it) — Lab 18."],
    ["Job / CronJob", "Run-to-completion work, and a scheduled Job (Lab 19)."],
    ["kubectl", "The Kubernetes command-line tool."],
])
p("You're done — congratulations! You've taken one app from a single container all the way "
  "to a scalable Kubernetes deployment.")

# ============================================================================
# RENDERERS
# ============================================================================
def _md_anchor(t):
    a = t.lower().replace("—", "").replace("(", "").replace(")", "")
    return "-".join(a.split()).replace("/", "").replace(".", "").replace(",", "").replace(":", "")

def _md_toc(blocks):
    out = ["## Table of Contents", ""]
    for b in blocks:
        if b[0] == "h2":
            out.append(f"- [{b[1]}](#{_md_anchor(b[1])})")
    out.append("")
    return "\n".join(out)

def _md_code_lang(c):
    if c.lstrip().startswith("FROM "): return "dockerfile"
    if "apiVersion:" in c or c.lstrip().startswith("services:"): return "yaml"
    return "bash"

def render_markdown(blocks):
    out, injected = [], False
    for b in blocks:
        k = b[0]
        if k == "h1":
            out.append(f"# {b[1]}\n")
            if not injected:
                out.append(f"**Course Code:** {CODE}  ·  **Version {VERSION}**  ·  Tertiary Infotech Academy Pte Ltd\n")
                out.append("### Document Version Control Record\n")
                out.append("| Version | Effective Date | Summary of Changes | Author |")
                out.append("| --- | --- | --- | --- |")
                for v in VERSIONS:
                    out.append(f"| {v[0]} | {v[1]} | {v[2]} | {v[3]} |")
                out.append("")
                out.append(_md_toc(blocks))
                injected = True
            continue
        elif k == "h2": out.append(f"## {b[1]}\n")
        elif k == "h3": out.append(f"### {b[1]}\n")
        elif k == "p": out.append(f"{b[1]}\n")
        elif k == "steps": out.append("\n".join(f"{i}. {s}" for i, s in enumerate(b[1], 1)) + "\n")
        elif k == "bullets": out.append("\n".join(f"- {s}" for s in b[1]) + "\n")
        elif k == "code": out.append(f"```{_md_code_lang(b[1])}\n{b[1]}\n```\n")
        elif k == "table":
            rows = b[1]
            cells = lambda r: "| " + " | ".join(c.replace("\n", "<br>") for c in r) + " |"
            out.append(cells(rows[0]))
            out.append("| " + " | ".join("---" for _ in rows[0]) + " |")
            for r in rows[1:]: out.append(cells(r))
            out.append("")
        elif k == "note": out.append(f"> **Note:** {b[1]}\n")
        elif k == "test": out.append(f"> ✅ **Test it:** {b[1]}\n")
        elif k == "rule": out.append("---\n")
    return "\n".join(out).strip() + "\n"

# ---- docx helpers ----
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc); tcPr.append(shd)

def _shade_para(pr, hexc="F3F5F8"):
    ppr = pr._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc); ppr.append(shd)

def _runs(paragraph, text):
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part: continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(part)

def render_docx(blocks):
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)
    prodoc.style_headings(doc)
    prodoc.add_cover_page(doc, "Learner Guide", TITLE, VERSION,
        org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
        course_logo=os.path.join(ASSETS, "docker-k8s-course-logo.png"))
    prodoc.add_version_control(doc, VERSIONS)
    prodoc.add_toc(doc, levels="1-2")
    for b in blocks:
        k = b[0]
        if k == "h1": continue
        elif k == "h2": doc.add_paragraph(style="Heading 1").add_run(b[1])
        elif k == "h3": doc.add_paragraph(style="Heading 2").add_run(b[1])
        elif k == "p": _runs(doc.add_paragraph(), b[1])
        elif k == "steps":
            for s in b[1]: _runs(doc.add_paragraph(style="List Number"), s)
        elif k == "bullets":
            for s in b[1]: _runs(doc.add_paragraph(style="List Bullet"), s)
        elif k == "code":
            pr = doc.add_paragraph(); _shade_para(pr)
            r = pr.add_run(b[1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif k == "table":
            rows = b[1]
            t = doc.add_table(rows=0, cols=len(rows[0])); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                cells = t.add_row().cells
                for ci, val in enumerate(row):
                    cells[ci].text = ""; pp = cells[ci].paragraphs[0]
                    if ri == 0:
                        rr = pp.add_run(val.replace("\n", " ")); rr.bold = True
                        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); rr.font.size = Pt(9.5)
                        _shade(cells[ci], "1F6FEB")
                    else:
                        for j, line in enumerate(val.split("\n")):
                            if j: pp = cells[ci].add_paragraph()
                            _runs(pp, line)
                            for rn in pp.runs: rn.font.size = Pt(9.5)
        elif k == "note":
            pr = doc.add_paragraph(); _shade_para(pr, "FFF4E5")
            rr = pr.add_run("Note:  "); rr.bold = True; rr.font.color.rgb = RGBColor(0xB5, 0x6A, 0x00)
            _runs(pr, b[1])
        elif k == "test":
            pr = doc.add_paragraph(); _shade_para(pr, "E8F7EE")
            rr = pr.add_run("✅ Test it:  "); rr.bold = True; rr.font.color.rgb = RGBColor(0x12, 0x7A, 0x3E)
            _runs(pr, b[1])
        elif k == "rule":
            pr = doc.add_paragraph()
            ppr = pr._p.get_or_add_pPr(); bdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "D0D7DE")
            bdr.append(bot); ppr.append(bdr)
    prodoc.add_page_numbers(doc)
    prodoc.enable_update_fields(doc)
    return doc

md = render_markdown(B)
with open(os.path.join(ROOT, "LG-Application-Integration-with-Docker-and-Kubernetes.md"), "w") as f:
    f.write(md)
render_docx(B).save(os.path.join(ROOT, "courseware",
                    "LG-Application-Integration-with-Docker-and-Kubernetes.docx"))
print("Wrote LG markdown + DOCX. sections(h2):", sum(1 for b in B if b[0] == "h2"),
      "| blocks:", len(B), "| md chars:", len(md))
