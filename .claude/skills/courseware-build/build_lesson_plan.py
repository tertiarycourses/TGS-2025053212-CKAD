#!/usr/bin/env python3
"""Generate the 2-day Lesson Plan (DOCX) for WSQ TGS-2021010366
'Application Integration with Docker and Kubernetes' from labs_data.py.

Daily schedule 9:00 AM - 6:00 PM = 8 training hours (1 h lunch, tea breaks within).
Day 2 final assessment starts at 4:00 PM. Every day's rows (excl. lunch) total 480 min."""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import prodoc
import labs_data as L
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.environ.get("COURSE_REPO") or os.getcwd()
TITLE = "Application Integration with Docker and Kubernetes"
CODE = "TGS-2021010366"
ASSETS = os.path.join(ROOT, "courseware", "assets")

VERSIONS = [
    ("1.0", "30 June 2026",
     "First version — 2-day lesson plan (Docker + Kubernetes) aligned to the slide deck "
     "and the 19 hands-on labs; 9:00 AM-6:00 PM, 8 training hours/day",
     "Tertiary Infotech Academy Pte Ltd"),
]

BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
HEADER_FILL, TOPIC_FILL, BREAK_FILL = "1F6FEB", "E8F0FE", "FFF4E5"

doc = Document()
doc.styles["Normal"].font.name = "Arial"; doc.styles["Normal"].font.size = Pt(11)

def set_cell_bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc); tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=10, align=None, italic=False):
    cell.text = ""; pp = cell.paragraphs[0]
    if align is not None: pp.alignment = align
    r = pp.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    return pp

def add_heading(text, space_before=12, space_after=6):
    pp = doc.add_paragraph(style="Heading 1")
    pp.paragraph_format.space_before = Pt(space_before); pp.paragraph_format.space_after = Pt(space_after)
    pp.add_run(text); return pp

# ---------------------------------------------------------------- front matter
prodoc.style_headings(doc)
prodoc.add_cover_page(doc, "LESSON PLAN", TITLE, "1.0",
    org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
    course_logo=os.path.join(ASSETS, "docker-k8s-course-logo.png"))
prodoc.add_version_control(doc, VERSIONS)
prodoc.add_toc(doc, levels="1-1")

info = [
    ("Course Title", TITLE),
    ("Course Code", CODE),
    ("Duration", "2 Days (16 training hours)"),
    ("Daily Schedule", "9:00 AM – 6:00 PM (8 training hours/day, excluding lunch)"),
    ("Lunch Break", "1:00 PM – 2:00 PM (1 hour)"),
    ("Breaks", "Short tea breaks are scheduled within each day's training hours"),
    ("Delivery Mode", "Instructor-led, hands-on labs in KillerCoda browser terminals"),
    ("Prerequisites", "Basic command-line familiarity; no prior Docker or Kubernetes experience required"),
    ("Tools", "Docker, Docker Compose, kubectl, KillerCoda, a Docker Hub account"),
]
tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in info:
    row = tbl.add_row().cells
    set_cell_text(row[0], k, bold=True, color=DARK, size=10); set_cell_bg(row[0], "F1F5FB")
    set_cell_text(row[1], v, size=10)
for row in tbl.rows:
    row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(4.6)

add_heading("Course Overview")
doc.add_paragraph(
    "This 2-day hands-on course teaches participants to containerise, run and orchestrate "
    "real applications with Docker and Kubernetes. Throughout the course learners build and "
    "deploy a single sample web application — TaskBoard (a Flask task tracker backed by Redis "
    "and PostgreSQL) — progressing from a single container to a multi-service Docker Compose "
    "stack and finally a scalable Kubernetes deployment. Day 1 covers Docker: commands, "
    "building images with a Dockerfile, CMD vs ENTRYPOINT, volumes, networking, configuration, "
    "Docker Hub and Docker Compose. Day 2 covers Kubernetes: Pods, Namespaces, Deployments, "
    "rolling updates and rollbacks, Services, storage, and Jobs/CronJobs.")

add_heading("Learning Outcomes")
doc.add_paragraph("By the end of this course, participants will be able to:")
for o in [
    "Explain containerisation and the Docker architecture, and manage containers with the Docker CLI.",
    "Build custom images with a Dockerfile, applying layer caching and .dockerignore best practices.",
    "Distinguish CMD from ENTRYPOINT and package command-line tools correctly.",
    "Persist and share data using named volumes and bind mounts.",
    "Connect containers over a custom network using DNS, and publish services with port mapping.",
    "Configure containers at runtime with environment variables and publish images to Docker Hub.",
    "Define and run multi-service applications (web + cache + database) with Docker Compose.",
    "Explain Kubernetes architecture and deploy applications as Pods, Deployments and Services.",
    "Scale and self-heal workloads and perform zero-downtime rolling updates and rollbacks.",
    "Persist data with PersistentVolumes/Claims and run batch work with Jobs and CronJobs.",
]:
    pp = doc.add_paragraph(o, style="List Bullet"); pp.paragraph_format.space_after = Pt(2)

# ---------------------------------------------------------------- schedule helpers
def schedule_table(rows):
    t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(["Time", "Topic / Activity", "Duration"]):
        set_cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_bg(hdr[i], HEADER_FILL)
    for time, activity, minutes, kind in rows:
        c = t.add_row().cells
        set_cell_text(c[0], time, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, bold=(kind == "topic"))
        set_cell_text(c[1], activity, size=9.5, bold=(kind in ("topic", "break")),
                      color=BRAND if kind == "topic" else (GREY if kind == "break" else None),
                      italic=(kind == "break"))
        set_cell_text(c[2], (f"{minutes} min" if minutes else "—"), size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill = TOPIC_FILL if kind == "topic" else (BREAK_FILL if kind == "break" else None)
        if fill:
            for cc in c: set_cell_bg(cc, fill)
    for row in t.rows:
        for i, w in enumerate([Inches(1.15), Inches(4.5), Inches(0.85)]): row.cells[i].width = w
    return t

def day_header(day, theme):
    add_heading(day, space_before=16)
    pp = doc.add_paragraph(); r = pp.add_run(theme); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY
    pp.paragraph_format.space_after = Pt(4)

def session_label(text):
    pp = doc.add_paragraph(); pp.paragraph_format.space_before = Pt(8); pp.paragraph_format.space_after = Pt(3)
    r = pp.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BRAND

def lab(n):
    l = L.lab_by_num(n)
    return f"Lab {n}: {l['title']}"

# ---------------------------------------------------------------- DAY 1 (Docker)
day_header("Day 1", "Docker — Images, Storage, Networking & Compose")
session_label("Morning Session  ·  9:00 AM – 1:00 PM")
schedule_table([
    ("9:00 – 9:15",  "Welcome, course objectives & overview (TGS-2021010366)", 15, "normal"),
    ("9:15 – 9:45",  "Topic 1: Containers, Docker architecture & core concepts", 30, "topic"),
    ("9:45 – 10:10", lab(1), 25, "normal"),
    ("10:10 – 10:35", lab(2), 25, "normal"),
    ("10:35 – 10:50", "Tea Break", 15, "break"),
    ("10:50 – 11:10", "Topic 2: Building images — Dockerfile, layers & build cache", 20, "topic"),
    ("11:10 – 11:45", lab(3), 35, "normal"),
    ("11:45 – 12:10", lab(4), 25, "normal"),
    ("12:10 – 12:45", lab(5), 35, "normal"),
    ("12:45 – 1:00",  "Morning Q&A", 15, "normal"),
])
schedule_table([("1:00 – 2:00", "Lunch Break", 60, "break")])
session_label("Afternoon Session  ·  2:00 PM – 6:00 PM")
schedule_table([
    ("2:00 – 2:20",  "Topic 3: Docker storage — volumes & bind mounts", 20, "topic"),
    ("2:20 – 2:50",  lab(6), 30, "normal"),
    ("2:50 – 3:10",  "Topic 4: Docker networking — bridges, DNS & ports", 20, "topic"),
    ("3:10 – 3:40",  lab(7), 30, "normal"),
    ("3:40 – 3:55",  "Tea Break", 15, "break"),
    ("3:55 – 4:15",  lab(8), 20, "normal"),
    ("4:15 – 4:35",  lab(9), 20, "normal"),
    ("4:35 – 4:50",  "Topic 5: Docker Compose — multi-service apps", 15, "topic"),
    ("4:50 – 5:10",  lab(10), 20, "normal"),
    ("5:10 – 5:30",  lab(11), 20, "normal"),
    ("5:30 – 5:55",  lab(12), 25, "normal"),
    ("5:55 – 6:00",  "Day 1 recap & wrap-up", 5, "normal"),
])

# ---------------------------------------------------------------- DAY 2 (Kubernetes)
day_header("Day 2", "Kubernetes — Deploy, Scale, Update & Persist")
session_label("Morning Session  ·  9:00 AM – 1:00 PM")
schedule_table([
    ("9:00 – 9:15",  "Day 1 recap & Day 2 objectives", 15, "normal"),
    ("9:15 – 9:55",  "Topic 6: Kubernetes architecture & kubectl", 40, "topic"),
    ("9:55 – 10:30", lab(13), 35, "normal"),
    ("10:30 – 10:55", lab(14), 25, "normal"),
    ("10:55 – 11:10", "Tea Break", 15, "break"),
    ("11:10 – 11:30", "Topic 7: Deployments, ReplicaSets & self-healing", 20, "topic"),
    ("11:30 – 12:10", lab(15), 40, "normal"),
    ("12:10 – 12:45", lab(16), 35, "normal"),
    ("12:45 – 1:00",  "Morning Q&A", 15, "normal"),
])
schedule_table([("1:00 – 2:00", "Lunch Break", 60, "break")])
session_label("Afternoon Session  ·  2:00 PM – 6:00 PM")
schedule_table([
    ("2:00 – 2:20",  "Topic 8: Services & networking in Kubernetes", 20, "topic"),
    ("2:20 – 2:55",  lab(17), 35, "normal"),
    ("2:55 – 3:10",  "Topic 9: Storage, Jobs & CronJobs", 15, "topic"),
    ("3:10 – 3:35",  lab(18), 25, "normal"),
    ("3:35 – 4:00",  lab(19), 25, "normal"),
    ("4:00 – 5:15",  "Final Assessment — written / MCQ on the LMS (lms.tertiaryinfotech.com)", 75, "topic"),
    ("5:15 – 5:30",  "Tea Break", 15, "break"),
    ("5:30 – 6:00",  "Course review, learner feedback (TRAQOM) & closing", 30, "normal"),
])

# ---------------------------------------------------------------- resources + assessment
add_heading("Labs Covered", space_before=16)
doc.add_paragraph("All 19 hands-on labs are delivered in KillerCoda browser terminals and "
                  "mirror the slide deck and Learner Guide exactly:")
rt = doc.add_table(rows=0, cols=3); rt.style = "Table Grid"; rt.alignment = WD_TABLE_ALIGNMENT.CENTER
hr = rt.add_row().cells
for i, h in enumerate(["Lab", "Topic", "Title"]):
    set_cell_text(hr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9.5); set_cell_bg(hr[i], HEADER_FILL)
for l in L.LABS:
    c = rt.add_row().cells
    set_cell_text(c[0], str(l["num"]), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(c[1], l["topic"], size=9)
    set_cell_text(c[2], l["title"], size=9)
for row in rt.rows:
    for i, w in enumerate([Inches(0.5), Inches(2.0), Inches(4.0)]): row.cells[i].width = w

add_heading("Assessment", space_before=14)
doc.add_paragraph(
    "Participants are assessed through their hands-on lab work across both days and a final "
    "written / MCQ assessment on Day 2 (starting at 4:00 PM) delivered on the LMS. A minimum "
    "of 75% attendance across all sessions is required for SSG funding.")

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)

OUT = os.path.join(ROOT, "courseware", "LP-Application-Integration-with-Docker-and-Kubernetes.docx")
doc.save(OUT)
print("Saved:", OUT)

# ---------------------------------------------------------------- sanity: 480 min/day
day1 = [15,30,25,25,15,20,35,25,35,15] + [20,30,20,30,15,20,20,15,20,20,25,5]
day2 = [15,40,35,25,15,20,40,35,15] + [20,35,15,25,25,75,15,30]
for name, mins in [("Day 1", day1), ("Day 2", day2)]:
    tot = sum(mins)
    print(f"{name}: {tot} min ({tot/60:.1f} h) ->", "OK" if tot == 480 else "MISMATCH")
